from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

log = open('write_config.log', 'w')

def Create_Code_Style(document):
    # Defining Document Styles for Code
    styles = document.styles
    style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['No Spacing']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(8)

def Create_Header1_Style(document):
    # Defining Document Styles for Code
    styles = document.styles
    style = styles.add_style('Header 1', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Heading 1']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

def Create_Header2_Style(document):
    # Defining Document Styles for Code
    styles = document.styles
    style = styles.add_style('Header 2', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Heading 2']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

def Create_Paragraph_Style(document):
    # Defining Document Styles for Code
    styles = document.styles
    style = styles.add_style('Paragraph 1', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Paragraph']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

def Write_Title(document, Title):
    document.add_heading(Title, 0).add_run().font.name = 'Calibri'
    #p = document.add_paragraph('A plain paragraph having some ')
    #p.add_run('bold').bold = True
    #p.add_run(' and some ')
    #p.add_run('italic.').italic = True

def Write_Body(document,Subtitle,Data, FileName):
    document.add_heading(Subtitle, level=1)
    document.add_paragraph(Data, style='Normal').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #document.add_page_break()
    document.save('%s.docx' % FileName)
    print >> log, ("Body Content has been written to docx file")
    print ("Body Content has been written to docx file")

def Write_Code(document, Subtitle, Code, FileName):
    document.add_heading(Subtitle, level=1)
    document.add_paragraph(Code, style='Code').alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_page_break()

    document.save('%s.docx' % FileName)
    print >> log, ("Code Content has been written to docx file")
    print ("Code Content has been written to docx file")

def main():
    document = Document()
    Data = ''
    title = 'Title of the document'
    subtitle = 'Item 1'
    #Data = "Lorem Ipsum is simply dummy text of the printing and typesetting industry. Lorem Ipsum has been the industry's standard dummy text ever since the 1500s, when an unknown printer took a galley of type and scrambled it to make a type specimen book. It has survived not only five centuries, but also the leap into electronic typesetting, remaining essentially unchanged. It was popularised in the 1960s with the release of Letraset sheets containing Lorem Ipsum passages, and more recently with desktop publishing software like Aldus PageMaker including versions of Lorem Ipsum."
    #Data = open("./File/CNS-ASW-KDN-402.txt", "r")
    with open("./File/CNS-ASW-KDN-402.txt", "r") as SourceText:
        for line in SourceText:
            Data += line.rstrip('\n')
    sub1 = 'Item 2'
    Data1 = "Lorem Ipsum is simply dummy text of the printing and typesetting industry. Lorem Ipsum has been the industry's standard dummy text ever since the 1500s, when an unknown printer took a galley of type and scrambled it to make a type specimen book. It has survived not only five centuries, but also the leap into electronic typesetting, remaining essentially unchanged. It was popularised in the 1960s with the release of Letraset sheets containing Lorem Ipsum passages, and more recently with desktop publishing software like Aldus PageMaker including versions of Lorem Ipsum."
    sub2 = 'Item 3'
    Data2 = "Lorem Ipsum is simply dummy text of the printing and typesetting industry. Lorem Ipsum has been the industry's standard dummy text ever since the 1500s, when an unknown printer took a galley of type and scrambled it to make a type specimen book. It has survived not only five centuries, but also the leap into electronic typesetting, remaining essentially unchanged. It was popularised in the 1960s with the release of Letraset sheets containing Lorem Ipsum passages, and more recently with desktop publishing software like Aldus PageMaker including versions of Lorem Ipsum."

    Write_Title(document,title)
    Write_Body(document, sub1, Data1, "test")
    Write_Body(document, sub2, Data2, "test")
    Write_Code(document, subtitle, Data, "test")
    #Save_Docx(document, "test")


if __name__ == "__main__":
    #If this Python file runs by itself, run below command. If imported, this section is not run
    main()